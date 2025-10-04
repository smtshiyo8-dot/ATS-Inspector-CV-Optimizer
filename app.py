"""
ATS-Inspector & CV Optimizer
single-file Flask app (app.py)

What this does
- Accept a vacancy link OR pasted job advert text
- Detect likely ATS/platform used by the employer (heuristic signatures)
- Ask for a CV upload (PDF, DOCX, TXT)
- Score the CV against the job advert and ATS heuristics (0-100)
- If score < 85: produce tailored improvement tips and an optional revamped DOCX that injects keywords and ATS-friendly structure

Requirements
- Python 3.9+
- pip install -r requirements.txt

requirements.txt (example contents)
Flask==2.2.5
requests==2.31.0
beautifulsoup4==4.12.2
python-docx==0.8.11
pdfplumber==0.7.8

(If pdfplumber gives trouble you can remove PDF support or use PyMuPDF)

Quick usage
1) Save this file as app.py in a folder (e.g., C:\projects\ats_inspector or ~/projects/ats_inspector)
2) Create and activate a virtual environment
   python -m venv .venv
   # Windows
   .\.venv\Scripts\activate
   # macOS / Linux
   source .venv/bin/activate
3) Create requirements.txt with the packages above
4) pip install -r requirements.txt
5) python app.py
6) Open http://127.0.0.1:5000 in your browser

Notes & disclaimers
- ATS detection uses heuristics (URL patterns + HTML signatures). It's accurate for many common ATS platforms but not guaranteed for custom career portals.
- This tool runs locally — no cloud AI required. Revamp is a keyword-driven restructure and not a semantic rewrite.
- Treat outputs as suggestions; always review the revamped CV before sending it to an employer.

"""

import os
import re
import math
import tempfile
from collections import Counter
from urllib.parse import urlparse

from flask import Flask, request, render_template_string, send_from_directory, redirect, url_for, flash
import requests
from bs4 import BeautifulSoup

# document parsing
import docx
import pdfplumber

# -------------------- Config --------------------
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.secret_key = os.environ.get('FLASK_SECRET', 'change-this-for-production')

# -------------------- ATS signatures --------------------
ATS_PATTERNS = {
    'Greenhouse': {
        'domains': ['greenhouse.io', 'greenhouse.io/people', 'greenhousehire.com'],
        'signatures': ['greenhouse', 'powered by greenhouse']
    },
    'Lever': {
        'domains': ['lever.co', 'jobs.lever.co', 'hire.lever.co'],
        'signatures': ['lever', 'apply.lever']
    },
    'Workday': {
        'domains': ['workday.com', 'myworkdayjobs.com', 'workdayjobs.com'],
        'signatures': ['workday']
    },
    'iCIMS': {
        'domains': ['icims.com', 'careers.icims.com'],
        'signatures': ['icims']
    },
    'SmartRecruiters': {
        'domains': ['smartrecruiters.com'],
        'signatures': ['smartrecruiters']
    },
    'Jobvite': {
        'domains': ['jobvite.com', 'jobs.jobvite.com'],
        'signatures': ['jobvite']
    },
    'Taleo': {
        'domains': ['taleo.net', 'oraclecloud.com', 'taleo.io'],
        'signatures': ['taleo']
    },
    'SuccessFactors': {
        'domains': ['successfactors.com', 'sapsf.com'],
        'signatures': ['successfactors', 'sap successfactors']
    },
    'Workable': {
        'domains': ['workable.com'],
        'signatures': ['workable']
    },
    'BambooHR': {
        'domains': ['bamboohr.com'],
        'signatures': ['bamboohr']
    },
    'ApplicantPro': {
        'domains': ['applicantpro.com'],
        'signatures': ['applicantpro']
    },
}

GENERIC_TIPS = [
    'Use a single-column layout: most ATS parse single-column text better than multi-column or tables.',
    'Prefer standard section headings: "Work Experience", "Education", "Skills", "Contact".',
    'Save as DOCX or a simple PDF (DOCX is usually safest). Avoid images, text in headers/footers, and fancy graphics.',
    'Use standard date formats (e.g., Apr 2020 – Feb 2023).',
    'Include a concise skills section listing keywords from the job advert.',
    'Avoid special characters and excessive symbols (they sometimes break parsers).',
]

# -------------------- Utilities --------------------

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def fetch_url_text(url):
    headers = {
        'User-Agent': 'ATS-Inspector/1.0 (+https://example.local)'
    }
    try:
        r = requests.get(url, headers=headers, timeout=10)
        r.raise_for_status()
        return r.text
    except Exception as e:
        print('fetch error', e)
        return None


# attempt to pull job description text from HTML using heuristics
def extract_job_text_from_html(html):
    soup = BeautifulSoup(html, 'html.parser')

    # common containers
    selectors = [
        "[class*=description]",
        "[class*=job-description]",
        "[class*=jobDescription]",
        "[id*=job-]",
        'article',
        'main',
        'section',
        '[class*=posting]'
    ]

    candidate_text = ''
    for sel in selectors:
        el = soup.select_one(sel)
        if el and el.get_text(strip=True):
            candidate_text = el.get_text(separator=' ', strip=True)
            if len(candidate_text) > 200:
                break

    # fallback: entire body
    if not candidate_text:
        body = soup.body
        if body:
            candidate_text = body.get_text(separator=' ', strip=True)

    # clean
    candidate_text = re.sub(r'\s+', ' ', candidate_text)
    return candidate_text


def detect_ats(url, html_text):
    url_l = (url or '').lower()
    html_l = (html_text or '').lower()

    # domain check
    parsed = urlparse(url_l) if url_l else None
    domain = parsed.netloc if parsed else ''

    matches = []

    for ats_name, data in ATS_PATTERNS.items():
        for d in data.get('domains', []):
            if d in url_l or d in domain:
                matches.append(ats_name)
                break
        for sig in data.get('signatures', []):
            if sig in html_l:
                matches.append(ats_name)
                break

    if matches:
        # return most common match
        c = Counter(matches).most_common(1)[0][0]
        return c

    # final heuristic: look for "powered by" strings
    m = re.search(r'powered by\s+([A-Za-z0-9\-]+)', html_l)
    if m:
        return m.group(1).capitalize()

    return 'Unknown/Custom'


# very small stopword list for keyword extraction
STOPWORDS = set("""
the a an and or to of for with on in at by from as is are that this will be
""".split())


def extract_keywords(text, top_n=25):
    # simple tokenization
    text = re.sub(r'[^A-Za-z0-9\-\s]', ' ', text or '')
    tokens = [t.lower() for t in text.split() if len(t) > 2]
    tokens = [t for t in tokens if t not in STOPWORDS]
    freqs = Counter(tokens)
    most = [w for w, _ in freqs.most_common(top_n)]
    return most


def extract_text_from_docx(path):
    doc = docx.Document(path)
    texts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text)
    return '\n'.join(texts)


def extract_text_from_pdf(path):
    text = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                ptext = page.extract_text()
                if ptext:
                    text.append(ptext)
    except Exception as e:
        print('pdfplumber error:', e)
    return '\n'.join(text)


def extract_text_from_txt(path):
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def extract_text_from_file(path):
    ext = path.rsplit('.', 1)[1].lower()
    if ext == 'docx':
        return extract_text_from_docx(path)
    elif ext == 'pdf':
        return extract_text_from_pdf(path)
    elif ext == 'txt':
        return extract_text_from_txt(path)
    else:
        return ''


# -------------------- Scoring --------------------

def find_contact_info(text):
    email = re.search(r'[\w\.-]+@[\w\.-]+', text or '')
    phone = re.search(r'(?:\+?\d{2,3}[\s-]?)?(?:\d{2,4}[\s-]?){2,4}\d{2,4}', text or '')
    return (email.group(0) if email else None, phone.group(0) if phone else None)


def score_resume(cv_text, job_keywords, job_title=None):
    # component weights
    weights = {
        'contact': 15,
        'title': 15,
        'keywords': 40,
        'length': 10,
        'formatting': 10
    }

    total = 0

    # contact
    email, phone = find_contact_info(cv_text)
    contact_score = 0
    if email and phone:
        contact_score = weights['contact']
    elif email or phone:
        contact_score = weights['contact'] * 0.6
    total += contact_score

    # title
    title_score = 0
    if job_title and job_title.lower() in (cv_text or '').lower():
        title_score = weights['title']
    total += title_score

    # keywords
    matched = 0
    job_keywords = [k.lower() for k in job_keywords]
    cv_lower = (cv_text or '').lower()
    if job_keywords:
        for k in set(job_keywords):
            if k in cv_lower:
                matched += 1
        keyword_score = weights['keywords'] * (matched / max(1, len(job_keywords)))
    else:
        keyword_score = weights['keywords'] * 0.5
    total += keyword_score

    # length
    length_score = 0
    chars = len(cv_text or '')
    if chars > 1000:
        length_score = weights['length']
    elif chars > 500:
        length_score = weights['length'] * 0.6
    total += length_score

    # formatting heuristics: heavy special characters, many tabs or very long lines -> penalty
    formatting_score = weights['formatting']
    non_ascii_ratio = sum(1 for c in cv_text or '' if ord(c) > 127) / max(1, len(cv_text or ''))
    tabs = (cv_text or '').count('\t')
    long_lines = sum(1 for line in (cv_text or '').splitlines() if len(line) > 200)
    if non_ascii_ratio > 0.02 or tabs > 5 or long_lines > 5:
        formatting_score *= 0.5
    total += formatting_score

    # normalize
    total = min(100, round(total))

    breakdown = {
        'contact_score': round(contact_score,1),
        'title_score': round(title_score,1),
        'keyword_score': round(keyword_score,1),
        'length_score': round(length_score,1),
        'formatting_score': round(formatting_score,1)
    }

    return total, breakdown


# -------------------- Revamp --------------------
from docx import Document


def create_revamped_docx(original_text, contact_email, contact_phone, job_title, job_keywords, out_path):
    doc = Document()

    # Header
    header_lines = []
    if contact_email:
        header_lines.append(contact_email)
    if contact_phone:
        header_lines.append(contact_phone)
    header_text = ' | '.join(header_lines) if header_lines else 'Contact information'
    doc.add_heading('Keyword-optimized CV', level=1)
    doc.add_paragraph(header_text)

    # Summary (keyword-injected)
    summary = 'Targeting role: {}'.format(job_title or 'Relevant role')
    if job_keywords:
        summary += '\nKey focus areas: ' + ', '.join(job_keywords[:12])
    doc.add_heading('Professional Summary', level=2)
    doc.add_paragraph(summary)

    # Skills
    doc.add_heading('Key Skills', level=2)
    if job_keywords:
        for k in job_keywords[:30]:
            doc.add_paragraph(k, style='List Bullet')
    else:
        doc.add_paragraph('See original CV for full skill details.')

    # Add original content under 'Full CV (for review)'
    doc.add_heading('Full CV (original content)', level=2)
    for line in (original_text or '').splitlines():
        if line.strip():
            # keep paragraphs manageable
            doc.add_paragraph(line.strip())

    doc.save(out_path)
    return out_path


# -------------------- Tips per ATS --------------------
ATS_TIPS = {
    'Greenhouse': [
        'Greenhouse parses plain DOCX well; avoid placing contact information inside headers/footers.',
        'Use clear bullets for accomplishments and avoid images or complex tables.'
    ],
    'Lever': [
        'Lever-based career pages prefer standard section headings and clear dates per job role.',
        'Avoid embedding text inside text-boxes or columns.'
    ],
    'Workday': [
        'Workday can struggle with multi-column layouts; use single column and plain fonts.',
        'Place your contact details at the top in plain text (not header/footer).' 
    ],
    'iCIMS': [
        'iCIMS works well with DOCX — keep skills in a dedicated section and use plain bullets.'
    ],
    'Taleo': [
        'Taleo-based portals sometimes struggle with special characters — remove emoji and fancy bullets.'
    ],
    'Unknown/Custom': [
        'This appears to be a custom portal. Follow generic ATS best-practices: simple DOCX, clear headings, no images.'
    ]
}


def build_improvement_tips(score, ats_name, cv_text, job_keywords):
    tips = []
    if score >= 85:
        tips.append('Great — your CV scored {:.0f}. Minor tweaks can still help, but it is ATS-friendly.'.format(score))
        return tips

    tips.append('Your CV scored {:.0f}/100. Here are prioritized improvement tips:'.format(score))

    # ATS-specific tips
    ats_specific = ATS_TIPS.get(ats_name, ATS_TIPS.get('Unknown/Custom'))
    tips.extend(ats_specific)

    # General tips (only add ones that seem relevant)
    if 'Work Experience' not in cv_text and 'Experience' not in cv_text:
        tips.append('Add a clear "Work Experience" section with job titles and dates.')

    email, phone = find_contact_info(cv_text)
    if not email:
        tips.append('Add a visible email address at the top of the CV.')
    if not phone:
        tips.append('Add a phone number at the top of the CV.')

    if job_keywords:
        # show top missing keywords
        cv_lower = (cv_text or '').lower()
        missing = [k for k in job_keywords if k.lower() not in cv_lower]
        if missing:
            tips.append('Include more of these job-specific keywords where truthful: {}'.format(', '.join(missing[:12])))

    tips.extend(GENERIC_TIPS[:3])

    return tips


# -------------------- Flask routes --------------------

INDEX_HTML = '''
<!doctype html>
<title>ATS-Inspector & CV Optimizer</title>
<h1>ATS-Inspector & CV Optimizer</h1>
<p>Paste a vacancy link or job advert, upload your CV, and get an ATS detection + CV score.</p>
<form method=post enctype=multipart/form-data action="/analyze">
  <label for="vacancy_link">Vacancy link (optional):</label><br>
  <input type=text name=vacancy_link style="width:90%"><br><br>
  <label for="job_text">Or paste the full job advert / description (optional):</label><br>
  <textarea name=job_text rows=8 style="width:90%"></textarea><br><br>
  <label for="file">Upload CV (DOCX / PDF / TXT):</label><br>
  <input type=file name=file><br><br>
  <input type=submit value='Analyze CV'>
</form>
<p>Note: this tool runs locally; detection uses heuristics and is not 100% guaranteed.</p>
'''

RESULT_HTML = '''
<!doctype html>
<title>Result - ATS Inspector</title>
<h1>Results</h1>
<p><strong>Detected ATS/Platform:</strong> {{ ats }}</p>
<p><strong>CV Score:</strong> {{ score }} / 100</p>
<h2>Breakdown</h2>
<ul>
  <li>Contact score: {{ breakdown.contact_score }}</li>
  <li>Title match score: {{ breakdown.title_score }}</li>
  <li>Keyword score: {{ breakdown.keyword_score }}</li>
  <li>Length score: {{ breakdown.length_score }}</li>
  <li>Formatting score: {{ breakdown.formatting_score }}</li>
</ul>

{% if tips %}
<h2>Improvement Tips</h2>
<ol>
{% for t in tips %}
  <li>{{ t }}</li>
{% endfor %}
</ol>
{% endif %}

{% if revamped_url %}
<p><a href="{{ revamped_url }}">Download the keyword-optimized (revamped) DOCX</a></p>
{% endif %}

<p><a href="/">Analyze another CV</a></p>
'''


@app.route('/', methods=['GET'])
def index():
    return render_template_string(INDEX_HTML)


@app.route('/analyze', methods=['POST'])
def analyze():
    vacancy_link = request.form.get('vacancy_link', '').strip()
    job_text = request.form.get('job_text', '').strip()
    uploaded = request.files.get('file')

    if not uploaded or uploaded.filename == '':
        flash('Please upload a CV file (DOCX, PDF or TXT).')
        return redirect(url_for('index'))

    if not allowed_file(uploaded.filename):
        flash('File type not allowed. Allowed: pdf, docx, txt')
        return redirect(url_for('index'))

    filename = secure_filename(uploaded.filename)
    save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    uploaded.save(save_path)

    # 1) get job advertisement text (from link or pasted)
    html_text = ''
    if vacancy_link:
        fetched = fetch_url_text(vacancy_link)
        if fetched:
            html_text = fetched
            if not job_text:
                job_text = extract_job_text_from_html(fetched)
    
    # 2) detect ATS
    ats = detect_ats(vacancy_link, html_text)

    # 3) extract job keywords
    job_keywords = extract_keywords(job_text or html_text or '')

    # 4) extract CV text
    cv_text = extract_text_from_file(save_path)

    # basic guessed job title - try to find a title in job_text
    job_title = None
    if job_text:
        # naive: take first line or first 6 words
        job_title = job_text.strip().split('\n')[0][:80]

    # 5) score
    score, breakdown = score_resume(cv_text, job_keywords, job_title)

    # 6) build tips
    tips = build_improvement_tips(score, ats if ats in ATS_TIPS else 'Unknown/Custom', cv_text, job_keywords)

    # 7) create revamped docx if score < 85
    revamped_url = None
    if score < 85:
        email, phone = find_contact_info(cv_text)
        # create safe output filename
        base = os.path.splitext(filename)[0]
        outname = secure_filename(base + '_revamped.docx')
        outpath = os.path.join(app.config['OUTPUT_FOLDER'], outname)
        try:
            create_revamped_docx(cv_text, email, phone, job_title, job_keywords, outpath)
            revamped_url = url_for('download_file', filename=outname)
        except Exception as e:
            print('revamp error', e)

    return render_template_string(RESULT_HTML, ats=ats, score=score, breakdown=breakdown, tips=tips, revamped_url=revamped_url)


from werkzeug.utils import secure_filename


@app.route('/outputs/<path:filename>')
def download_file(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
