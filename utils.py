from io import BytesIO

def detect_ats(file):
    # Fake ATS detection for demo
    return "Yes"  # Or "No"

def score_cv(file, job_description):
    # Fake scoring for demo
    score = 85
    improvements = [
        "Add more relevant skills.",
        "Use action verbs in experience section."
    ]
    # Extract keywords from job description (simplified)
    keywords = job_description.lower().split()[:10]  # take first 10 words as demo
    return score, improvements, keywords

def modify_cv_docx(file, keywords):
    from docx import Document
    doc = Document()
    # Standard CV headings
    doc.add_heading('Full Name', level=1)
    doc.add_paragraph('Professional Summary: Experienced professional skilled in ' + ', '.join(keywords) + '.')
    doc.add_heading('Experience', level=2)
    doc.add_paragraph('Include experience details here and relevant keywords like ' + ', '.join(keywords))
    doc.add_heading('Education', level=2)
    doc.add_paragraph('Include your education here.')
    doc.add_heading('Skills', level=2)
    doc.add_paragraph(', '.join(keywords))
    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def modify_cv_pdf(file, keywords):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from io import BytesIO

    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Full Name")
    y -= 30
    c.setFont("Helvetica", 12)
    c.drawString(50, y, "Professional Summary: Experienced professional skilled in " + ', '.join(keywords))
    y -= 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Experience")
    y -= 25
    c.setFont("Helvetica", 12)
    c.drawString(50, y, "Include experience details here incorporating keywords like " + ', '.join(keywords))
    y -= 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Education")
    y -= 25
    c.setFont("Helvetica", 12)
    c.drawString(50, y, "Include your education here.")
    y -= 30
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Skills")
    y -= 25
    c.setFont("Helvetica", 12)
    c.drawString(50, y, ', '.join(keywords))
    c.showPage()
    c.save()
    bio.seek(0)
    return bio
